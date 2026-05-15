from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_mediqueue_doc():
    doc = Document()

    # Title
    title = doc.add_heading('MediQueue Project Documentation', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. Project Name
    doc.add_heading('1. Project Name', level=1)
    doc.add_paragraph('MediQueue (Intelligent Medical Queue & Appointment Management System)')

    # 2. Project Idea
    doc.add_heading('2. Project Idea', level=1)
    doc.add_paragraph(
        'MediQueue is an integrated digital platform designed specifically for the healthcare sector. '
        'It aims to transform clinic management from traditional manual methods to a smart digital approach '
        'by providing tools for online appointment booking and real-time queue tracking. '
        'This reduces physical waiting time in clinics and enhances operational efficiency.'
    )

    # 3. Problem Statement
    doc.add_heading('3. Problem Statement', level=1)
    p = doc.add_paragraph()
    p.add_run('The project addresses several challenges faced by patients and doctors:').bold = True
    doc.add_paragraph('Waiting Room Overcrowding: Leading to poor patient experience and increased risk of infection.', style='List Bullet')
    doc.add_paragraph('Appointment Inaccuracy: Traditional appointments are often delayed without the patient knowledge, causing frustration.', style='List Bullet')
    doc.add_paragraph('Burdensome Manual Management: Reliance on paper records or simple spreadsheets leads to booking errors and scheduling conflicts.', style='List Bullet')
    doc.add_paragraph('Accessibility Issues: Patients often need to call or visit the clinic just to check available slots.', style='List Bullet')

    # 4. Project Goals
    doc.add_heading('4. Project Goals', level=1)
    doc.add_paragraph('Queue Organization: Provide a live tracking system that allows patients to know their number in the queue and expected entry time from home.', style='List Bullet')
    doc.add_paragraph('Enhanced Patient Experience: Minimize actual in-clinic waiting time to the lowest possible levels.', style='List Bullet')
    doc.add_paragraph('Increased Doctor Efficiency: Provide a dashboard that enables doctors to manage their schedules and cases easily.', style='List Bullet')
    doc.add_paragraph('Transparency: Make information about clinics, specialties, and available appointments clear and accessible to everyone.', style='List Bullet')

    # 5. Project Plan / Timeline
    doc.add_heading('5. Project Plan / Timeline', level=1)
    doc.add_paragraph('The project follows a structured 4-month timeline:')
    doc.add_picture('/mnt/ahmed/MediQueue/gantt.png', width=Inches(6))
    doc.add_paragraph('Month 1: Planning & Design (Requirements Gathering, Database Design, UI/UX Design).', style='List Number')
    doc.add_paragraph('Month 2-3: Development (Core Framework, Business Logic, Frontend Integration).', style='List Number')
    doc.add_paragraph('Month 4: Testing & Deployment (System Testing, Bug Fixing, Beta Launch).', style='List Number')

    # 6. Users / Actors
    doc.add_heading('6. Users / Actors', level=1)
    doc.add_paragraph('Patient: The person who searches for a clinic, makes a booking, and tracks their turn.', style='List Bullet')
    doc.add_paragraph('Doctor: Responsible for providing the service and managing the patient queue.', style='List Bullet')
    doc.add_paragraph('Admin: Responsible for overall system management, clinics, and user accounts.', style='List Bullet')

    # 7. Functional Requirements
    doc.add_heading('7. Functional Requirements', level=1)
    doc.add_paragraph('Patient Features: Search, Book, Cancel, Live Status.', style='List Bullet')
    doc.add_paragraph('Doctor Features: Set Availability, Manage Queue (Call Next, End Session).', style='List Bullet')
    doc.add_paragraph('Admin Features: User Management, Clinic/Doctor Management, Statistics.', style='List Bullet')

    # 8. Non-Functional Requirements
    doc.add_heading('8. Non-Functional Requirements', level=1)
    doc.add_paragraph('Security: Data protection via Identity system and encryption.', style='List Bullet')
    doc.add_paragraph('Performance: Fast search and real-time queue updates.', style='List Bullet')
    doc.add_paragraph('Availability: 24/7 system uptime for bookings.', style='List Bullet')

    # 9. Use Case Diagram
    doc.add_heading('9. Use Case Diagram', level=1)
    doc.add_picture('/mnt/ahmed/MediQueue/use_case.png', width=Inches(5))
    doc.add_paragraph(
        'The system coordinates actions between three main actors. Patients interact with booking and tracking services, '
        'Doctors manage the live session and availability, while Admins handle the infrastructure and user database.'
    )

    # --- Arabic Section ---
    doc.add_page_break()
    
    title_ar = doc.add_heading('توثيق مشروع MediQueue', 0)
    title_ar.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 1. اسم المشروع
    doc.add_heading('1. اسم المشروع', level=1)
    doc.add_paragraph('MediQueue (نظام إدارة الطوابير والمواعيد الطبية الذكي)')

    # 2. فكرة المشروع
    doc.add_heading('2. فكرة المشروع', level=1)
    doc.add_paragraph(
        'مشروع MediQueue هو منصة رقمية متكاملة مصممة خصيصاً للقطاع الطبي. '
        'تهدف المنصة إلى تحويل عملية إدارة العيادات من الأسلوب اليدوي التقليدي إلى أسلوب رقمي ذكي، '
        'من خلال توفير أدوات لحجز المواعيد عبر الإنترنت، ومتابعة طوابير الانتظار بشكل مباشر.'
    )

    # 3. المشكلة التي يحلها المشروع
    doc.add_heading('3. المشكلة التي يحلها المشروع', level=1)
    doc.add_paragraph('يعالج المشروع عدة تحديات تواجه المرضى والأطباء:')
    doc.add_paragraph('الازدحام في غرف الانتظار.', style='List Bullet')
    doc.add_paragraph('عدم دقة المواعيد وصعوبة متابعتها.', style='List Bullet')
    doc.add_paragraph('الإدارة اليدوية المرهقة والأخطاء الناتجة عنها.', style='List Bullet')
    doc.add_paragraph('صعوبة الوصول للمواعيد المتاحة.', style='List Bullet')

    # 4. أهداف المشروع
    doc.add_heading('4. أهداف المشروع', level=1)
    doc.add_paragraph('تنظيم الطوابير بشكل رقمي حي.', style='List Bullet')
    doc.add_paragraph('تحسين تجربة المريض وتقليل وقت الانتظار.', style='List Bullet')
    doc.add_paragraph('رفع كفاءة الطبيب في إدارة المواعيد.', style='List Bullet')
    doc.add_paragraph('الشفافية في عرض المواعيد المتاحة.', style='List Bullet')

    # 5. خطة المشروع
    doc.add_heading('5. خطة المشروع (Timeline)', level=1)
    doc.add_paragraph('المرحلة الأولى: التخطيط والتصميم (شهر واحد).', style='List Number')
    doc.add_paragraph('المرحلة الثانية: التطوير والبرمجة (شهرين).', style='List Number')
    doc.add_paragraph('المرحلة الثالثة: الاختبار والإطلاق (شهر واحد).', style='List Number')

    # 6. المستخدمون
    doc.add_heading('6. المستخدمون', level=1)
    doc.add_paragraph('المريض: البحث والحجز والمتابعة.', style='List Bullet')
    doc.add_paragraph('الطبيب: إدارة التوفر والطابور الحي.', style='List Bullet')
    doc.add_paragraph('المدير: إدارة النظام والعيادات والأطباء.', style='List Bullet')

    doc.save('/mnt/ahmed/MediQueue/doc.docx')
    print("Document saved successfully to /mnt/ahmed/MediQueue/doc.docx")

if __name__ == "__main__":
    create_mediqueue_doc()
